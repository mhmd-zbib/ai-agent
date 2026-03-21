from io import BytesIO

from app.modules.pipeline.parsers.base import BaseParser, ParsedDocument, clean_text


class DocxParser(BaseParser):
    """
    Extracts text from .docx files using python-docx.

    Heading paragraphs are preserved as-is; blank paragraphs are dropped.
    No page concept — ``pages`` is always ``None``.
    """

    def parse(self, content: bytes, file_name: str = "") -> ParsedDocument:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(BytesIO(content))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                paragraphs.append(stripped)

        text = clean_text("\n\n".join(paragraphs))
        return ParsedDocument(text=text, pages=None)
